# !pip install pyspellchecker
# !pip install textblob
# !pip install python-docx PyPDF2 langdetect language-tool-python

import re
from typing import List, Dict
from spellchecker import SpellChecker
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import docx
import PyPDF2
from langdetect import detect
import language_tool_python
import json
from datetime import datetime
from typing import List
import os
import json
import pandas as pd
import language_tool_python

# Инициализация инструмента проверки орфографии
language_tool = language_tool_python.LanguageTool('ru')

# Чтение текста из DOCX файла
def read_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Чтение текста из PDF файла
def read_pdf(file_path: str) -> str:
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

# Загрузка требований из файла JSON
def load_requirements(file_path: str) -> List[str]:
    with open(file_path, 'r', encoding='utf-8') as file:
        requirements = json.load(file)
    return requirements.get("job_keywords", [])

# Проверка орфографии и грамматики
def check_spelling_and_grammar(text: str) -> Dict[str, List[str]]:
    technology_keywords = ["Java", "Python", "SQL", "C++", "JavaScript", "HTML", "CSS"]
    matches = language_tool.check(text)
    errors = []
    error_words = []

    for match in matches:
        error_fragment = text[match.offset:match.offset + match.errorLength]
        # Пропускаем ключевые слова технологий
        if not any(tech.lower() in error_fragment.lower() for tech in technology_keywords):
            errors.append(f"Ошибка: {match.message}, Позиция: {match.offsetInContext}")
            error_words.append(error_fragment.strip())  # Добавляем само слово с ошибкой

    return {
        "errors": errors,
        "error_words": list(set(error_words))  # Убираем дублирования слов
    }

# Проверка ключевых слов
def check_keywords(resume: str, job_keywords: List[str]) -> Dict[str, List[str]]:
    found_keywords = [kw for kw in job_keywords if kw.lower() in resume.lower()]
    missing_keywords = list(set(job_keywords) - set(found_keywords))
    return {
        "found_keywords": found_keywords,
        "missing_keywords": missing_keywords,
    }

# Проверка хронологии
def check_timing(resume: str) -> List[str]:
    date_range_pattern = re.compile(
        r'(?P<start>\b\d{1,2}\s+[А-яёЁ]+\s+\d{4}|\b[А-яёЁ]+\s+\d{4}|\d{4})\s*-\s*(?P<end>\b\d{1,2}\s+[А-яёЁ]+\s+\d{4}|\b[А-яёЁ]+\s+\d{4}|\d{4})'
    )
    months = {
        "январь": 1, "февраль": 2, "март": 3, "апрель": 4, "май": 5, "июнь": 6,
        "июль": 7, "август": 8, "сентябрь": 9, "октябрь": 10, "ноябрь": 11, "декабрь": 12
    }

    def parse_date(date_str: str) -> datetime:
        try:
            date_str = date_str.strip()
            if re.match(r'\d{1,2}\s+[А-яёЁ]+\s+\d{4}', date_str):
                day, month_text, year = date_str.split()
                return datetime(int(year), months[month_text.lower()], int(day))
            elif re.match(r'[А-яёЁ]+\s+\d{4}', date_str):
                month_text, year = date_str.split()
                return datetime(int(year), months[month_text.lower()], 1)
            elif re.match(r'\d{4}', date_str):
                return datetime(int(date_str), 1, 1)
        except KeyError:
            raise ValueError(f"Ошибка: Неизвестный месяц '{month_text}' в дате '{date_str}'")
        except Exception as e:
            raise ValueError(f"Ошибка при обработке даты '{date_str}': {e}")

    periods = []
    for match in date_range_pattern.finditer(resume):
        start_date = parse_date(match.group('start'))
        end_date = parse_date(match.group('end'))
        periods.append((start_date, end_date))

    periods.sort(key=lambda x: x[0])

    issues = []
    for i in range(len(periods) - 1):
        current_end = periods[i][1]
        next_start = periods[i + 1][0]
        gap = (next_start - current_end).days
        if gap > 183:
            issues.append(
                f"Обнаружен перерыв в хронологии более полугода между {current_end.strftime('%d.%m.%Y')} "
                f"и {next_start.strftime('%d.%m.%Y')}."
            )

    return issues

# Генерация суммаризации
def generate_summary(skills: List[str], job_keywords: List[str]) -> str:
    combined_skills = list(set(skills + job_keywords))
    summary = f"Владею следующими технологиями: {', '.join(combined_skills)}."
    return summary

# Вставка суммаризации после телефона
def insert_summary_into_corrected_text(corrected_text: str, summary: str) -> str:
    phone_pattern = re.compile(r"(Телефон[:\-]?\s*\+?\d[\d\s\-\(\)]*)", re.IGNORECASE)
    match = phone_pattern.search(corrected_text)

    if match:
        phone_section = match.group(0)
        summary_text = f"\n\n{summary}\n"
        return corrected_text.replace(phone_section, phone_section + summary_text, 1)
    return corrected_text + f"\n\n{summary}"

# Формирование улучшенного текста
def generate_corrected_text(text: str) -> str:
    matches = language_tool.check(text)
    corrected_text = text
    technology_keywords = ["Java", "Python", "SQL", "C++", "JavaScript", "HTML", "CSS"]
    for match in matches:
        error_fragment = text[match.offset:match.offset + match.errorLength]
        if not any(tech.lower() in error_fragment.lower() for tech in technology_keywords):
            corrected_text = language_tool.correct(corrected_text)
    return corrected_text

# Анализ резюме
def analyze_resume(file_path: str, job_keywords: List[str]) -> Dict:
    if file_path.endswith('.docx'):
        text = read_docx(file_path)
    elif file_path.endswith('.pdf'):
        text = read_pdf(file_path)
    else:
        raise ValueError("Формат файла не поддерживается. Используйте DOCX или PDF.")

    if detect(text) != 'ru':
        raise ValueError("Язык резюме должен быть русским.")

    spelling_and_grammar_errors = check_spelling_and_grammar(text)
    keyword_check = check_keywords(text, job_keywords)
    timing_issues = check_timing(text)

    skills_section = re.search(r"Навыки:\s*(.*)", text, re.IGNORECASE)
    skills = [skill.strip() for skill in skills_section.group(1).split(",")] if skills_section else []

    corrected_text = generate_corrected_text(text)
    summary = generate_summary(skills, job_keywords)
    corrected_text_with_summary = insert_summary_into_corrected_text(corrected_text, summary)

    match_score = len(keyword_check["found_keywords"]) / len(job_keywords) * 100 if job_keywords else 0
    return {
        "original_text": text,
        "corrected_text": corrected_text_with_summary,
        "summary": summary,
        "spelling_and_grammar_errors": spelling_and_grammar_errors,
        "missing_keywords": keyword_check["missing_keywords"],
        "timing_issues": timing_issues,
        "match_score": match_score
    }

# Функция для сохранения исправленного резюме
def save_resume_to_word(resume_text: str, file_name: str):
    output_dir = r"C:\Users\Athur\resume-analyzer"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    file_path = os.path.join(output_dir, file_name)  # Убедимся, что используется правильный путь
    doc = docx.Document()
    for line in resume_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(file_path)

    print(f"Сгенерированное резюме сохранено по пути: {file_path}")

# Основной блок кода
if __name__ == "__main__":
    # Указываем путь к файлу резюме
    print("Загружается файл резюме (DOCX или PDF):")
    uploaded = {"C:\\Users\\Athur\\resume-analyzer\\resume.docx": "resume.docx"}  # Симулируем загрузку файла
    resume_file_path = list(uploaded.keys())[0]

    # Указываем путь к файлу с требованиями
    print("Загружается файл с требованиями (JSON):")
    uploaded_requirements = {"C:\\Users\\Athur\\resume-analyzer\\notebooks\\requirements.json": "requirements.json"}  # Симулируем загрузку файла
    requirements_file_path = list(uploaded_requirements.keys())[0]

    # Загрузка ключевых слов из JSON
    with open(requirements_file_path, 'r', encoding='utf-8') as file:
        job_keywords = json.load(file).get("job_keywords", [])

    # Анализ резюме
    try:
        result = analyze_resume(resume_file_path, job_keywords)

        print("\n--- Исправленное резюме ---\n")
        print(result["corrected_text"])

        # Сохранение исправленного резюме
        save_resume_to_word(result["corrected_text"], "Исправленное_резюме.docx")

        print("\n--- Замечания ---\n")
        print("Ошибки орфографии и грамматики:")
        if result["spelling_and_grammar_errors"]["errors"]:
            print("\n".join(result["spelling_and_grammar_errors"]["errors"]))
        else:
            print("Ошибок не найдено.")

        print("\nПропущенные ключевые слова:")
        print(", ".join(result["missing_keywords"]) if result["missing_keywords"] else "Все ключевые слова найдены.")

        print(f"\nСоответствие требованиям: {result['match_score']:.2f}%")
    except ValueError as e:
        print(f"Ошибка: {e}")
