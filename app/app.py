import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from parser import parse_vacancies, technologies

# Функция для извлечения текста из PDF
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Функция для извлечения текста из DOCX
def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# Функция для определения категории резюме
def determine_resume_category(text):
    if "студент" in text.lower() or "учёба" in text.lower():
        return "student"
    elif "опыт работы" in text.lower() or "профессионал" in text.lower():
        return "professional"
    return "general"

# Заглушка для рекомендаций
def generate_recommendations(category):
    if category == "student":
        return [
            "Добавьте портфолио ваших учебных проектов.",
            "Укажите курсы и сертификаты, которые вы прошли.",
            "Сфокусируйтесь на навыках и потенциале."
        ]
    elif category == "professional":
        return [
            "Укажите конкретные достижения в цифрах (например, увеличил продажи на 20%).",
            "Добавьте ссылки на ваши работы или проекты.",
            "Сделайте акцент на ключевых навыках, связанных с вакансией."
        ]
    else:
        return [
            "Сократите объём резюме до 1-2 страниц.",
            "Проверьте текст на наличие орфографических ошибок.",
            "Добавьте контактные данные, если их нет."
        ]

# Функция для создания PDF
def generate_pdf(recommendations):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    c.setFont("Helvetica", 12)
    c.drawString(100, 750, "Рекомендации по улучшению резюме:")

    y_position = 730
    for rec in recommendations:
        c.drawString(100, y_position, f"- {rec}")
        y_position -= 20

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer

# Функция для создания DOCX
def generate_docx(recommendations):
    doc = Document()
    doc.add_heading('Рекомендации по улучшению резюме', 0)

    for rec in recommendations:
        doc.add_paragraph(f"- {rec}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Прогресс-бар для обработки файла
def file_processing_progress_bar():
    progress_bar = st.progress(0)
    for i in range(100):
        progress_bar.progress(i + 1)

# Интерактивные настройки для рекомендаций
def interactive_recommendations():
    st.sidebar.title("Настройка рекомендаций")
    recommendation_type = st.sidebar.radio("Выберите тип рекомендаций", ["Для студентов", "Для профессионалов", "Общие"])
    return recommendation_type

# Функция для сравнения текста резюме с навыками из вакансий
def compare_resume_with_vacancies(resume_text, tech_summary):
    missing_technologies = []
    for tech, _ in tech_summary:
        if tech.lower() not in resume_text.lower():
            missing_technologies.append(tech)
    return missing_technologies

# Заголовок приложения
st.title("Анализатор Резюме и Вакансий", anchor="top")

# Боковая панель
st.sidebar.title("Настройки")
st.sidebar.write("Выберите параметры для анализа:")
color_coding = st.sidebar.checkbox("Цветовая кодировка рекомендаций", value=True)

# Поле для загрузки файла
uploaded_file = st.file_uploader("Загрузите файл резюме (PDF или DOCX)", type=["pdf", "docx"])

# Поле для ввода профессии
vacancy_name = st.sidebar.text_input("Введите название профессии для анализа", value="Data Scientist")

# Обработка загруженного файла
if uploaded_file is not None:
    st.success(f"Файл '{uploaded_file.name}' успешно загружен!")

    # Прогресс-бар
    file_processing_progress_bar()

    file_text = ""

    if uploaded_file.name.endswith(".pdf"):
        file_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        file_text = extract_text_from_docx(uploaded_file)

    if file_text:
        # Определение категории резюме
        category = determine_resume_category(file_text)
        st.subheader("Категория резюме")
        st.write(f"Определена категория: **{category.capitalize()}**")

        # Генерация рекомендаций
        st.subheader("Рекомендации по улучшению резюме")

        # Интерактивная настройка рекомендаций
        recommendation_type = interactive_recommendations()
        if recommendation_type == "Для студентов":
            recommendations = generate_recommendations("student")
        elif recommendation_type == "Для профессионалов":
            recommendations = generate_recommendations("professional")
        else:
            recommendations = generate_recommendations("general")

        # Вывод рекомендаций с цветовой кодировкой
        for rec in recommendations:
            if color_coding:
                st.markdown(f"<p style='color: #4CAF50; font-size: 16px;'>{rec}</p>", unsafe_allow_html=True)
            else:
                st.markdown(f"<p style='font-size: 16px;'>{rec}</p>", unsafe_allow_html=True)

        # Кнопка для скачивания PDF
        pdf_file = generate_pdf(recommendations)
        st.download_button("Скачать рекомендации (PDF)", pdf_file, file_name="recommendations.pdf", mime="application/pdf")

        # Кнопка для скачивания DOCX
        docx_file = generate_docx(recommendations)
        st.download_button("Скачать рекомендации (DOCX)", docx_file, file_name="recommendations.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Парсинг вакансий
        if st.button("Анализировать вакансии"):
            with st.spinner("Сбор данных о вакансиях..."):
                df, tech_summary = parse_vacancies(vacancy_name, technologies, days=30)

            st.success("Парсинг завершён!")
            st.subheader("Топ-10 технологий из вакансий")
            for tech, count in tech_summary:
                st.write(f"{tech.capitalize()}: {count} упоминаний")

            # Сравнение навыков из резюме с вакансиями
            missing_skills = compare_resume_with_vacancies(file_text, tech_summary)

            if missing_skills:
                st.warning("В резюме отсутствуют следующие навыки:")
                for skill in missing_skills:
                    st.write(f"- {skill}")
            else:
                st.success("Ваше резюме полностью соответствует требованиям!")

            # Генерация рекомендаций по вакансиям
            recommendations_from_vacancies = [f"Добавьте в резюме навык: {skill}" for skill in missing_skills]

            # Кнопка для скачивания рекомендаций на основе вакансий
            pdf_file_vacancies = generate_pdf(recommendations_from_vacancies)
            st.download_button("Скачать рекомендации по вакансиям (PDF)", pdf_file_vacancies, file_name="recommendations_vacancies.pdf", mime="application/pdf")

    else:
        st.error("Не удалось извлечь текст из файла. Проверьте его содержимое.")
